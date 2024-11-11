from flask import Flask, request, jsonify, render_template, send_from_directory
import fitz  # PyMuPDF for PDF handling
from langdetect import detect
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from datetime import datetime
import os
from werkzeug.utils import secure_filename
import shutil
import re
import requests
import json
from dotenv import load_dotenv
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)

load_dotenv()

api_key = os.getenv('API_KEY')
APY_KEY = api_key

# Configura a pasta temporária
app.config['UPLOAD_FOLDER'] = os.path.join(app.root_path, 'uploads')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Função para reiniciar a pasta de uploads
def restart():
    uploads_path = os.path.join(app.root_path, 'uploads')
    if os.path.exists(uploads_path):
        shutil.rmtree(uploads_path)
        os.makedirs(uploads_path, exist_ok=True)
    return None

# Função para extrair texto de um arquivo
def extract_text(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    if extension == '.pdf':
        return pdf_to_text(file_path)
    elif extension == '.docx':
        return docx_to_text(file_path)
    elif extension == '.txt':
        return txt_to_text(file_path)
    else:
        return None

# Função para extrair texto de PDF
def pdf_to_text(file_path):
    text = ""
    with fitz.open(file_path) as pdf:
        for page in pdf:
            text += page.get_text()
    return text

# Função para extrair texto de DOCX
def docx_to_text(file_path):
    doc = Document(file_path)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return text

# Função para extrair texto de TXT
def txt_to_text(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return text

# Função para verificar se o arquivo tem mais de 2 páginas
def has_more_than_two_pages(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    if extension == '.pdf':
        with fitz.open(file_path) as pdf:
            return pdf.page_count > 2
    elif extension == '.docx':
        doc = Document(file_path)
        return len(doc.paragraphs) > 75
    elif extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            return len(text) > 5000
    else:
        return False

# Função para verificar se o arquivo contém imagem
def contains_image(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    if extension == '.pdf':
        with fitz.open(file_path) as pdf:
            for page_num in range(pdf.page_count):
                page = pdf.load_page(page_num)
                if len(page.get_images(full=True)) > 0:
                    return True
        return False
    elif extension == '.docx':
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                drawing_element = run._element.xml
                drawing_element=str(drawing_element)
                flagTest= '<w:drawing>' in drawing_element
                if flagTest:
                    return True
        return False
    else:
        return False

# Função para verificar se o texto está em português
def is_portuguese(text):
    try:
        return detect(text) == 'pt'
    except:
        return False

# Função para verificar se o texto contém informações pessoais
def check_personalinfo(text):
    phone_pattern = r'\s*\(\d{2}\)\s*\d{5}-\d{4}'
    email_pattern = r'\s*[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    linkedin_pattern= r'(?:\/\/(?:www|m)\.)?(?:linkedin\.com|linked-in\.co\.uk)(?:\/in\/)[a-zA-Z0-9-]+'
    github_pattern = r'(?:\/\/(?:www\.)?github\.com|github\.io)\/[A-Za-z0-9-]+'
    phone = re.search(phone_pattern, text)
    email = re.search(email_pattern, text)
    linkedin = re.search(linkedin_pattern, text)
    github = re.search(github_pattern, text)
    if (phone is not None) or (email is not None) or (linkedin is not None) or (github is not None) :
        return True 
    else:
        return False 

# Função para verificar erros de português
def check_portuguese_errors(text):
    text=str(text).upper()
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent"
    headers = {"Content-Type": "application/json"}
    data = {
        "contents": [
            {"parts": [{"text": f"Analise o seguinte texto em português e me retorne em formato JSON uma lista de termos que possuem **erros de ortografia graves**, incluindo sugestões de correção. Considere **erros de ortografia graves** que prejudicam a legibilidade e a compreensão do texto, como:- Erros de ortografia (ex: \\\"presedente\\\" - errado); - Erros de acentuação (ex: \\\"facil\\\" - errado); - Erros de hífen (ex: \\\"mal-humorado\\\" - errado). **Importante:** O texto de entrada está todo em maiúsculas. **Ignore qualquer análise de maiúsculas e minúsculas e não corrija a ortografia de palavras em maiúsculas.** Ignore erros de estilo, abreviações, nomenclatura, erros de gramática e erros de pontuação. Caso o texto não tenha erros de ortografia graves, retorne uma lista vazia. \n\n**Texto:** {text}\n\n**Formato JSON:** {{'errors': [{{'termo': 'termo_com_erro', 'correcao':'correção1', 'mensagem': 'Mensagem sobre o erro.'}}]}}"}]}
                ]
            } 
    
    response = requests.post(url, headers=headers, json=data, params={"key": APY_KEY})
    if response.status_code == 200:
        try:
            result = response.json()
            results = result["candidates"][0]["content"]['parts'][0]["text"]
            results = str(results).replace("```json", "")
            results = results.replace("```", "")
            result = json.loads(results)
            return result["errors"]
        except (KeyError, IndexError, json.JSONDecodeError):
            return []
    else:
        return []

# Função para remover strings irrelevantes
def remove_strings_useless(input_string):
    input_string=str(input_string).upper()
    return input_string.replace(" ", "")

# Função para validar os erros retornados pelo Gemini
def validateReturnGemini(errors):
    valid_items = []
    unique_items = set()
    
    for error in errors:
        if remove_strings_useless(error['termo']) != remove_strings_useless(error['correcao']) and tuple(error.items()) not in unique_items:
            valid_items.append(error)
            unique_items.add(tuple(error.items()))
    return valid_items

# Função para verificar palavras-chave
def check_keywords(job_description):
    job_description=str(job_description).upper()
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent"
    headers = {"Content-Type": "application/json"}
    data = {
        "contents": [
            {"parts": [{"text": f"Extraia as palavras-chave relevantes de um **Texto:** de descrição de vaga, focando em habilidades técnicas, comportamentais e experiências profissionais. As palavras-chave devem se referir a:* **Habilidades técnicas:**  Linguagens de programação, frameworks, ferramentas, softwares, plataformas.* **Habilidades comportamentais:**  Adjetivos que descrevem características desejáveis como comunicação, trabalho em equipe, proatividade, criatividade etc.* **Experiência profissional:**  Verbos que indicam ações realizadas como 'desenvolver', 'implementar', 'integrar', 'configurar', 'gerenciar', etc.Apresente as palavras-chave em formato JSON, com a chave 'palavras_chaves' e uma lista com as palavras-chave.**Texto:** {job_description}**Formato JSON:** {{'palavras_chaves': []}}"}]}
        ]
    }
    response = requests.post(url, headers=headers, json=data, params={"key": APY_KEY})
    if response.status_code == 200:
        try:
            result = response.json()
            results = result["candidates"][0]["content"]['parts'][0]["text"]
            results = str(results).replace("```json", "")
            results = results.replace("```", "")
            result = json.loads(results)
            return result["palavras_chaves"]
        except (KeyError, IndexError, json.JSONDecodeError):
            return []
    else:
        return []

# Função para comparar palavras-chave do currículo com a descrição da vaga
def match_keywords_with_resume(job_description, keywords):
    keywords=list(keywords)
    counter=0
    keywords_missing=""
    list_keywords_missing=[]
    uppered_list=list()
    job_description=str(job_description).upper()
    for keyword in keywords:
        uppered_list.append(str(keyword).upper())
    for keyword in uppered_list:
        if keyword in job_description:
            counter=counter+1
        else:
            list_keywords_missing.append(keyword)
    
    if len(keywords)==0:
        score = 0
    else:
        score=counter/len(keywords)
    if len(list_keywords_missing)!=0:
        keywords_missing = ', '.join(list_keywords_missing)
        if len(list_keywords_missing)==1:
            keywords_missing=keywords_missing[:-1]
    return score, keywords_missing

# Função para calcular a similaridade entre o currículo e a descrição da vaga
def calculate_similarity(resume_text, job_description):
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform([resume_text, job_description])
    similarity_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    return similarity_score

# Rota principal
@app.route('/')
def home():
    current_year = datetime.now().year
    return render_template('index.html', current_year=current_year)

# Rota para a página de dicas
@app.route('/dicas')
def dicas():
    current_year = datetime.now().year
    return render_template('dicas.html', current_year=current_year)

# Rota para a página de validação do CV
@app.route('/validar')
def validar():
    current_year = datetime.now().year
    return render_template('validar.html', current_year=current_year)

# Rota para a página de modelos
@app.route('/modelo')
def modelo():
    current_year = datetime.now().year
    return render_template('modelo.html', current_year=current_year)

# Rotas para download dos modelos
@app.route('/download/<nome_arquivo>')
def download_modelo(nome_arquivo):
    return send_from_directory('static/modelos_CV', nome_arquivo, as_attachment=True)

# Rota para upload do currículo
@app.route('/upload', methods=['POST'])
def upload_resume():
    if 'file' not in request.files:
        return jsonify({"error": "File missing"}), 400
    
    file = request.files['file']
    job_description = request.form.get('job_description')

    # Salva o arquivo na pasta temporária
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    # Realiza a análise
    result = analyze_resume(file_path, job_description)

    # Limpa a pasta de uploads
    restart()

    # Retorna o resultado para o cliente
    return jsonify(result[0]), result[1]

# Função para analisar o currículo
def analyze_resume(file_path, job_description):
    resume_text = extract_text(file_path)

    if resume_text is None:
        return {'result': 'Rejected', 'reason': 'Invalid file format. Please upload a PDF, DOCX, or TXT file.'}, 400

    if has_more_than_two_pages(file_path):
        return {'result': 'Rejected', 'reason': 'Resume has more than 2 pages'}, 400

    if contains_image(file_path):
        return {'result': 'Rejected', 'reason': 'Resume contains a photo'}, 400

    if not is_portuguese(resume_text):
        return {'result': 'Rejected', 'reason': 'Resume is not in Portuguese'}, 400

    if not check_personalinfo(resume_text):
        return {'result': 'Rejected', 'reason': 'Resume Has no Email or Phone Number or Linkedin or Github for contact'}, 400

    errors = check_portuguese_errors(resume_text)
    if len(errors)!=0:
        errors=validateReturnGemini(errors)
        if len(errors)!=0:
            return {'result': 'Rejected', 'reason': 'Resume contains Portuguese grammar errors', 'errors': errors}, 400

    keywords = check_keywords(job_description)
    score, keywords_missing = match_keywords_with_resume(job_description, keywords)
    similarity_score = calculate_similarity(resume_text, job_description)

    return {
        "result": "Accepted",
        "reason": "Resume meets all criteria",
        "similarity_score": round(similarity_score * 100, 2),
        "contextual_score": round(score * 100, 2),
        "keywords_missing": keywords_missing
    }, 200

# Iniciar o servidor
if __name__ == '__main__':
    app.run(debug=True)